
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Error: 'python-docx' is not installed.")
    print("Please run: pip install python-docx")
    sys.exit(1)

def create_documentation_docx():
    # Paths
    script_dir = Path(__file__).parent
    project_root = script_dir.parent.parent
    md_path = project_root / "docs" / "PROJECT_MASTER_MANUAL.md"
    output_path = project_root / "Career_Copilot_RAG_Documentation.docx"

    if not md_path.exists():
        print(f"Error: Usage manual not found at {md_path}")
        return

    # Create Document
    doc = Document()
    
    # Title Style
    style_title = doc.styles['Title']
    style_title.font.name = 'Arial'
    style_title.font.size = Pt(26)
    style_title.font.bold = True
    style_title.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue

    # Heading 1
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Arial'
    style_h1.font.size = Pt(20)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(46, 116, 181) # Blue

    # Heading 2
    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Arial'
    style_h2.font.size = Pt(16)
    style_h2.font.bold = True
    style_h2.font.color.rgb = RGBColor(68, 114, 196) # Light Blue

    # Heading 3
    style_h3 = doc.styles['Heading 3']
    style_h3.font.name = 'Arial'
    style_h3.font.size = Pt(14)
    style_h3.font.bold = True
    style_h3.font.color.rgb = RGBColor(0, 0, 0)

    # Parsing
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_buffer = []

    for line in lines:
        stripped = line.strip()
        
        # Code Block Handling
        if stripped.startswith("```"):
            if in_code_block:
                # End of block -> Write buffer
                p = doc.add_paragraph()
                runner = p.add_run("\n".join(code_buffer))
                runner.font.name = 'Courier New'
                runner.font.size = Pt(9)
                p.style = 'No Spacing'
                
                # Add a border or background is hard in python-docx simple usage, 
                # so we stick to font change
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_buffer.append(line.rstrip())
            continue

        # Skip empty lines (doc adds spacing naturally)
        if not stripped:
            continue

        # Headers
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        
        # Bullets
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, stripped[2:])

        # Numbered List (Simple heuristic)
        elif re.match(r"^\d+\.", stripped):
            # Split "1. text"
            text_part = re.sub(r"^\d+\.\s*", "", stripped)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, text_part)
            
        # Blockquote
        elif stripped.startswith("> "):
            p = doc.add_paragraph(style='Quote')
            _add_formatted_text(p, stripped[2:])

        # Separator
        elif stripped.startswith("---"):
            doc.add_paragraph("_" * 50)

        # Normal Text
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

    doc.save(str(output_path))
    print(f"✅ Successfully created: {output_path}")

def _add_formatted_text(paragraph, text):
    """Parses **bold** and `code` spans."""
    # Split by bold markers
    # This is a basic parser. It handles **bold** and `code`.
    
    # We iterate chunks
    # Note: Regex splitting is tricky for nested, so we do a simple sequential scan
    # or just split by one type. Let's support **Bold** primarily.
    
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle `code` inside non-bold parts
            subparts = re.split(r"(`.*?`)", part)
            for sub in subparts:
                if sub.startswith("`") and sub.endswith("`"):
                    run = paragraph.add_run(sub[1:-1])
                    run.font.name = 'Courier New'
                    run.font.color.rgb = RGBColor(220, 20, 60) # Red-ish for code
                else:
                    paragraph.add_run(sub)

if __name__ == "__main__":
    create_documentation_docx()
